from ..config import get_settings
from typing import List, Optional, Dict
import logging
import os
import re

logger = logging.getLogger(__name__)
settings = get_settings()


class TextProcessor:
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        self.supported_extensions = {'.txt', '.pdf', '.doc', '.docx', '.xlsx', '.xls', '.csv', '.md'}

    def read_file(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext not in self.supported_extensions:
            raise ValueError(f"不支持的文件格式: {file_ext}，支持的格式: {', '.join(self.supported_extensions)}")

        readers = {
            '.txt': self._read_txt,
            '.pdf': self._read_pdf,
            '.doc': self._read_docx,
            '.docx': self._read_docx,
            '.xlsx': self._read_excel,
            '.xls': self._read_excel,
            '.csv': self._read_csv,
            '.md': self._read_markdown,
        }

        reader = readers.get(file_ext, self._read_txt)
        content = reader(file_path)
        return self.clean_text(content)

    def _read_txt(self, file_path: str) -> str:
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"无法识别文件编码: {file_path}")

    def _read_pdf(self, file_path: str) -> str:
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_parts = []
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[第{i+1}页]\n{page_text}")
                return '\n\n'.join(text_parts)
        except ImportError:
            raise ImportError("处理PDF需要安装PyPDF2: pip install PyPDF2")

    def _read_docx(self, file_path: str) -> str:
        try:
            import docx
            doc = docx.Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(f"\n[表格]\n{table_text}\n")

            return '\n\n'.join(text_parts)
        except ImportError:
            raise ImportError("处理DOCX需要安装python-docx: pip install python-docx")

    def _extract_table_text(self, table) -> str:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(' | '.join(cells))
        return '\n'.join(rows)

    def _read_excel(self, file_path: str) -> str:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            text_parts = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_text = f"[工作表: {sheet_name}]\n"
                rows = []

                for row in ws.iter_rows(values_only=True):
                    cells = [str(cell) if cell is not None else '' for cell in row]
                    if any(cell.strip() for cell in cells):
                        rows.append(' | '.join(cells))

                if rows:
                    sheet_text += '\n'.join(rows)
                    text_parts.append(sheet_text)

            wb.close()
            return '\n\n'.join(text_parts)
        except ImportError:
            raise ImportError("处理Excel需要安装openpyxl: pip install openpyxl")

    def _read_csv(self, file_path: str) -> str:
        import csv
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, newline='') as f:
                    reader = csv.reader(f)
                    rows = []
                    for row in reader:
                        if any(cell.strip() for cell in row):
                            rows.append(' | '.join(cell.strip() for cell in row))
                    return '\n'.join(rows)
            except UnicodeDecodeError:
                continue
        raise ValueError(f"无法识别CSV文件编码: {file_path}")

    def _read_markdown(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content

    def split_text(self, text: str, metadata: Optional[Dict] = None) -> List[str]:
        if not text:
            return []

        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = min(start + self.chunk_size, text_len)

            if end < text_len:
                best_break = end
                for sep in ['\n\n', '\n', '。', '！', '？', '；', '.', '!', '?', ';', '，', ',']:
                    last_sep = text[start:end].rfind(sep)
                    if last_sep > self.chunk_size // 3:
                        best_break = start + last_sep + len(sep)
                        break
                end = best_break

            chunk = text[start:end].strip()
            if chunk and len(chunk) > 10:
                chunk_metadata = {
                    'start_pos': start,
                    'end_pos': end,
                }
                if metadata:
                    chunk_metadata.update(metadata)
                chunks.append(chunk)

            start = end - self.chunk_overlap
            if start >= text_len or start < 0:
                break

        logger.info(f"文本分割完成: {len(chunks)} 个文本块")
        return chunks

    def clean_text(self, text: str) -> str:
        if not text:
            return ''

        text = text.replace('\r\n', '\n').replace('\r', '\n')
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)

        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            stripped = line.strip()
            if stripped:
                cleaned_lines.append(stripped)

        return '\n'.join(cleaned_lines)

    def extract_keywords(self, text: str, top_k: int = 10) -> List[str]:
        import jieba.analyse
        try:
            keywords = jieba.analyse.extract_tags(text, topK=top_k, withWeight=False)
            return keywords
        except ImportError:
            logger.warning("jieba未安装，无法提取关键词")
            return []

    def get_file_info(self, file_path: str) -> Dict:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        stat = os.stat(file_path)
        file_ext = os.path.splitext(file_path)[1].lower()

        return {
            'name': os.path.basename(file_path),
            'extension': file_ext,
            'size': stat.st_size,
            'size_display': self.format_size(stat.st_size),
            'modified_time': stat.st_mtime,
        }

    @staticmethod
    def format_size(size_bytes: int) -> str:
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        elif size_bytes < 1024 * 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
        else:
            return f"{size_bytes / (1024 * 1024 * 1024):.2f} GB"
